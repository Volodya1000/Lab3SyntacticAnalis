import os
from natasha import Doc, Segmenter, MorphVocab, NewsEmbedding, NewsMorphTagger, NewsSyntaxParser
import tkinter as tk
from tkinter import filedialog, messagebox
import window_utils as wu
from docx import Document
import sys
import time
from text_units import TextSegment
from token_model import LexicalToken
import lang_config as lc

class TextAnalyzerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("DocAnalyzer 1.4")
        self.root.geometry("1280x720")
        wu.align_window_center(self.root, 1280, 720)
        self.current_text = ""
        self.analysis_output = ""
        self.work_directory = os.getcwd()

        # NLP компоненты
        self.segmenter = Segmenter()
        self.morph_vocab = MorphVocab()
        self.embedding = NewsEmbedding()
        self.tagger = NewsMorphTagger(self.embedding)
        self.parser = NewsSyntaxParser(self.embedding)

        self.setup_interface()
        self.create_menu()

    def create_menu(self):
        menu_bar = tk.Menu(self.root)
        
        # Меню Файл
        file_menu = tk.Menu(menu_bar, tearoff=0)
        file_menu.add_command(label="Выбрать рабочую папку", command=self.select_folder)
        file_menu.add_command(label="Экспорт результатов", command=self.save_results)
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=sys.exit)
        menu_bar.add_cascade(label="Файл", menu=file_menu)
        
        # Меню Анализ
        analysis_menu = tk.Menu(menu_bar, tearoff=0)
        analysis_menu.add_command(label="Обработать документ", command=self.process_full_text)
        analysis_menu.add_command(label="Анализ выделения", command=self.process_selected_text)
        menu_bar.add_cascade(label="Действия", menu=analysis_menu)

        help_menu = tk.Menu(menu_bar, tearoff=0)
        help_menu.add_command(label="Как пользоваться", command=self.show_help)
        menu_bar.add_cascade(label="Помощь", menu=help_menu)


        self.root.config(menu=menu_bar)

    def setup_interface(self):
        self.dir_label = tk.Label(self.root, text=self.work_directory)
        self.dir_label.pack(pady=5)

        content_frame = tk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        content_frame.pack(fill=tk.BOTH, expand=True)

        self.file_selector = tk.Listbox(content_frame)
        content_frame.add(self.file_selector)
        self.file_selector.bind("<Double-Button-1>", self.open_file)
        self.refresh_file_list()

        self.text_panel = tk.Text(content_frame, wrap=tk.WORD)
        content_frame.add(self.text_panel)

        self.result_panel = tk.Text(content_frame, wrap=tk.WORD)
        content_frame.add(self.result_panel)

    def select_folder(self):
        new_dir = filedialog.askdirectory()
        if new_dir:
            self.work_directory = new_dir
            self.dir_label.config(text=self.work_directory)
            self.refresh_file_list()

    def refresh_file_list(self):
        self.file_selector.delete(0, tk.END)
        for item in os.listdir(self.work_directory):
            if item.endswith(".docx"):
                self.file_selector.insert(tk.END, item)

    def open_file(self, event):
        selected = self.file_selector.get(self.file_selector.curselection())
        full_path = os.path.join(self.work_directory, selected)
        try:
            doc = Document(full_path)
            self.current_text = "\n".join([para.text for para in doc.paragraphs])
            self.text_panel.delete(1.0, tk.END)
            self.text_panel.insert(tk.END, self.current_text)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать файл: {str(e)}")

    def process_full_text(self):
        content = self.text_panel.get(1.0, tk.END).strip()
        self.analyze_content(content)

    def process_selected_text(self):
        try:
            content = self.text_panel.get(tk.SEL_FIRST, tk.SEL_LAST).strip()
        except tk.TclError:
            messagebox.showerror("Ошибка", "Не выделен текст для анализа")
            return
        self.analyze_content(content)

  

    def analyze_content(self, text):
        start_time = time.time()  # Засекаем время начала обработки
        
        doc = Doc(text)
        doc.segment(self.segmenter)
        doc.tag_morph(self.tagger)
        doc.parse_syntax(self.parser)

        report = []
        for idx, sent in enumerate(doc.sents, 1):
            segment = TextSegment(sent.text)
            report.append(f"Предложение #{idx}:")
            
            for token in sent.tokens:
                if token.pos == "PUNCT":
                    continue
                
                token.lemmatize(self.morph_vocab)
                lemma = token.lemma or "-"
                pos = lc.PARTS_OF_SPEECH.get(token.pos, token.pos)
                rel = lc.SYNTACTIC_ROLES.get(token.rel, token.rel)
                
                lex_token = LexicalToken(token.text, lemma, pos, rel)
                segment.add_token(lex_token)
                report.append(
                    f"    токен: {lex_token.form}, "
                    f"лемма: {lex_token.base}, "
                    f"часть речи: {lex_token.category}, "
                    f"роль: {lex_token.function}"
                )
            report.append("")

        self.analysis_output = "\n".join(report)
        self.result_panel.delete(1.0, tk.END)
        self.result_panel.insert(tk.END, self.analysis_output)
        
        end_time = time.time()  # Засекаем время окончания обработки
        elapsed_time = end_time - start_time  # Вычисляем затраченное время
        
        # Показываем сообщение с временем обработки
        messagebox.showinfo(
            "Анализ завершен",
            f"Текст успешно обработан.\nВремя обработки: {elapsed_time:.2f} секунд"
        )

    def save_results(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not path:
            return
        
        try:
            doc = Document()
            
            # Добавляем заголовок
            doc.add_heading('Результаты анализа', 1)
            
            # Форматируем содержимое
            for line in self.analysis_output.split('\n'):
                if line.startswith('Сегмент #'):
                    doc.add_heading(line, level=2)
                elif line.startswith('    Лексема:'):
                    parts = line.split(', ')
                    p = doc.add_paragraph()
                    runner = p.add_run(parts[0] + '\n')
                    runner.bold = True
                    for part in parts[1:]:
                        p.add_run(part + '\n')
                else:
                    doc.add_paragraph(line)
            
            doc.save(path)
            messagebox.showinfo("Экспорт завершен", 
                f"Документ успешно сохранен:\n{path}")
                
        except Exception as e:
            messagebox.showerror("Ошибка экспорта", 
                f"Не удалось сохранить файл:\n{str(e)}")
            
    def show_help(self):
        help_text = """DocAnalyzer 1.4 - Руководство пользователя

    1. Выбор рабочей папки:
    - Используйте меню 'Файл > Выбрать рабочую папку'
    - Программа будет отображать все .docx файлы в выбранной папке

    2. Открытие документа:
    - Дважды кликните на файле в списке слева
    - Текст документа появится в центральной панели

    3. Анализ текста:
    - 'Обработать документ' - анализирует весь текст
    - 'Анализ выделения' - анализирует только выделенный фрагмент

    4. Результаты:
    - Результаты анализа отображаются в правой панели
    - Для каждого предложения выводится список токенов с:
        * Оригинальной формой
        * Нормальной формой (леммой)
        * Частью речи
        * Синтаксической ролью

    5. Экспорт:
    - Результаты можно сохранить в .docx через меню 'Файл > Экспорт результатов'
    """
        help_window = tk.Toplevel(self.root)
        help_window.title("Справка")
        help_window.geometry("600x400")
        wu.align_window_center(help_window, 600, 400)
        
        text_widget = tk.Text(help_window, wrap=tk.WORD, padx=10, pady=10)
        text_widget.insert(tk.END, help_text)
        text_widget.config(state=tk.DISABLED)
        text_widget.pack(fill=tk.BOTH, expand=True)
        
        close_button = tk.Button(help_window, text="Закрыть", command=help_window.destroy)
        close_button.pack(pady=5)